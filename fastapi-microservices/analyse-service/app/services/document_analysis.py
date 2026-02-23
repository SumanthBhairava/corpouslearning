import os
import re
import hashlib
import logging
from dataclasses import dataclass
from typing import Any, Dict, List, Optional
from xml.etree import ElementTree as ET

import requests
from docx import Document

from .lm_studio_client import LMStudioClient

logger = logging.getLogger(__name__)


@dataclass
class RedlineContext:
    redline_text: str
    redline_type: str
    addition_text: str = ""
    author: str = "Unknown"
    surrounding_context: str = ""
    comments: Optional[List[str]] = None
    position: int = 0


class AnalyseEngine:
    def __init__(
        self,
        injection_service_url: str,
        lm_studio_url: str,
        model_name: str,
        request_timeout: int = 30,
    ):
        self.injection_service_url = injection_service_url.rstrip("/")
        self.request_timeout = request_timeout
        self.llm = LMStudioClient(base_url=lm_studio_url, model=model_name, timeout=max(request_timeout, 45))

    def analyze_file(
        self,
        tenant_id: str,
        file_path: str,
        top_k: int,
        version_mode: str,
        contract_type: Optional[str],
    ) -> Dict[str, Any]:
        extraction = self._extract_document(file_path)
        doc_text = extraction["content"]

        contract = self._determine_contract_type(os.path.basename(file_path), doc_text)
        if contract_type:
            contract = contract_type

        # Always compare the uploaded document against precedent index first,
        # even if tracked changes cannot be extracted.
        baseline_query = self._build_baseline_query(doc_text, extraction["comments"])
        baseline_retrieval = self._retrieve_precedents_safe(
            tenant_id=tenant_id,
            query=baseline_query,
            top_k=top_k,
            version_mode=version_mode,
            contract_type=contract,
        )

        redline_insights = []
        for idx, change in enumerate(extraction["tracked_changes"][:30], start=1):
            context = self._build_surrounding_context(doc_text, change.get("text", ""))
            redline = RedlineContext(
                redline_text=change.get("text", ""),
                redline_type=change.get("type", "UNKNOWN"),
                surrounding_context=context,
                comments=extraction["comments"][:8],
                position=idx,
            )

            section_analysis = self._analyze_redline(redline)
            queries = self._generate_search_queries(section_analysis)
            best_query = self._pick_best_query(queries, redline.redline_text)
            retrieval = self._retrieve_precedents_safe(
                tenant_id=tenant_id,
                query=best_query,
                top_k=top_k,
                version_mode=version_mode,
                contract_type=contract,
            )

            negotiation = self._extract_negotiation_intelligence(
                redline=redline,
                section_analysis=section_analysis,
                retrieval_results=retrieval.get("results", []),
            )

            redline_insights.append(
                {
                    "index": idx,
                    "redline": {
                        "type": redline.redline_type,
                        "text": redline.redline_text,
                    },
                    "section_analysis": section_analysis,
                    "search_queries": queries,
                    "retrieval_total": retrieval.get("total", 0),
                    "citations": retrieval.get("results", []),
                    "negotiation_intelligence": negotiation,
                }
            )

        return {
            "file_name": os.path.basename(file_path),
            "file_hash": self._md5(file_path),
            "contract_type": contract,
            "document_precedent_retrieval": baseline_retrieval,
            "document_stats": {
                "characters": len(doc_text),
                "comments": len(extraction["comments"]),
                "tracked_changes": len(extraction["tracked_changes"]),
                "extraction_confidence": extraction["extraction_confidence"],
            },
            "redline_insights": redline_insights,
            "summary": {
                "redlines_analyzed": len(redline_insights),
                "comments_found": len(extraction["comments"]),
                "additions": len([c for c in extraction["tracked_changes"] if c.get("type") in ("ADDITION", "UNDERLINE")]),
                "strikeouts": len([c for c in extraction["tracked_changes"] if c.get("type") in ("STRIKEOUT", "DELETION")]),
                "document_level_precedents": int(baseline_retrieval.get("total", 0)),
            },
        }

    def _extract_document(self, file_path: str) -> Dict[str, Any]:
        doc = Document(file_path)
        content = "\n".join([p.text.strip() for p in doc.paragraphs if (p.text or "").strip()])

        tracked_changes = []
        confidence = 0.3

        try:
            root = ET.fromstring(doc.element.xml)
            for el in root.iter():
                if el.tag.endswith("}ins"):
                    txt = self._collect_text(el)
                    if txt:
                        tracked_changes.append({"type": "ADDITION", "text": txt})
                elif el.tag.endswith("}del"):
                    txt = self._collect_text(el)
                    if txt:
                        tracked_changes.append({"type": "DELETION", "text": txt})
            if tracked_changes:
                confidence = 0.9
        except Exception:
            logger.exception("XML extraction failed for tracked changes")

        if not tracked_changes:
            for para in doc.paragraphs:
                for run in para.runs:
                    txt = (run.text or "").strip()
                    if not txt:
                        continue
                    if run.font.strike:
                        tracked_changes.append({"type": "STRIKEOUT", "text": txt})
                    elif run.underline:
                        tracked_changes.append({"type": "UNDERLINE", "text": txt})
            if tracked_changes:
                confidence = 0.5

        comments = []
        markers = ["comment:", "note:", "rationale:", "requested:", "proposed:", "discussed"]
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if text and any(m in text.lower() for m in markers):
                comments.append(text)

        return {
            "content": content,
            "tracked_changes": tracked_changes,
            "comments": comments,
            "extraction_confidence": confidence,
        }

    def _collect_text(self, xml_node: ET.Element) -> str:
        parts = []
        for item in xml_node.iter():
            if item.tag.endswith("}t") or item.tag.endswith("}delText"):
                t = (item.text or "").strip()
                if t:
                    parts.append(t)
        return " ".join(parts)

    def _determine_contract_type(self, filename: str, content: str) -> str:
        prompt = (
            "Identify contract type from filename and content. Return JSON: "
            '{"contract_type":"...","confidence":"high|medium|low"}.\n\n'
            "FILENAME: %s\nCONTENT:\n%s"
        ) % (filename, content[:1800])
        response = self.llm.query(
            prompt=prompt,
            system_message="You are a legal contract classifier. Return JSON only.",
            temperature=0.0,
            max_tokens=180,
        )
        if response.success:
            try:
                data = self.llm.parse_json(response.content)
                return data.get("contract_type", "UNKNOWN")
            except Exception:
                pass
        return "UNKNOWN"

    def _analyze_redline(self, redline: RedlineContext) -> Dict[str, Any]:
        prompt = (
            "Analyze this redline change and return JSON with keys: section_concept, "
            "section_title, legal_topic, change_type, change_magnitude, "
            "negotiation_intent, search_strategy.\n\n"
            "REDLINE TYPE: %s\nREDLINE TEXT: %s\nCONTEXT:\n%s\nCOMMENTS:\n%s"
        ) % (
            redline.redline_type,
            redline.redline_text[:1000],
            redline.surrounding_context[:1200],
            "\n".join(redline.comments or [])[:600],
        )
        response = self.llm.query(
            prompt=prompt,
            system_message="You are an expert legal analyst. Return strict JSON only.",
            temperature=0.1,
            max_tokens=700,
        )
        if not response.success:
            return {
                "section_concept": "Unknown",
                "section_title": "Unknown",
                "legal_topic": "Unknown",
                "change_type": "unknown",
                "change_magnitude": "unknown",
                "negotiation_intent": "unknown",
                "search_strategy": "use redline text directly",
            }
        try:
            return self.llm.parse_json(response.content)
        except Exception:
            return {
                "section_concept": "Unknown",
                "section_title": "Unknown",
                "legal_topic": "Unknown",
                "change_type": "unknown",
                "change_magnitude": "unknown",
                "negotiation_intent": "unknown",
                "search_strategy": "use redline text directly",
            }

    def _generate_search_queries(self, section_analysis: Dict[str, Any]) -> Dict[str, List[str]]:
        prompt = (
            "Generate precedent search queries and return JSON keys: primary_queries, "
            "secondary_queries, comment_searches, section_indicators.\n\n"
            "ANALYSIS: %s"
        ) % json_safe(section_analysis)

        response = self.llm.query(
            prompt=prompt,
            system_message="You are a legal precedent researcher. Return JSON only.",
            temperature=0.2,
            max_tokens=500,
        )
        if not response.success:
            return {
                "primary_queries": [section_analysis.get("section_concept", "")],
                "secondary_queries": [section_analysis.get("legal_topic", "")],
                "comment_searches": [],
                "section_indicators": [],
            }
        try:
            return self.llm.parse_json(response.content)
        except Exception:
            return {
                "primary_queries": [section_analysis.get("section_concept", "")],
                "secondary_queries": [section_analysis.get("legal_topic", "")],
                "comment_searches": [],
                "section_indicators": [],
            }

    def _retrieve_precedents(
        self,
        tenant_id: str,
        query: str,
        top_k: int,
        version_mode: str,
        contract_type: Optional[str],
    ) -> Dict[str, Any]:
        body = {
            "tenant_id": tenant_id,
            "query": query,
            "top_k": top_k,
            "version_mode": version_mode,
        }
        if contract_type and contract_type != "UNKNOWN":
            body["contract_type"] = contract_type

        response = requests.post(
            self.injection_service_url + "/rag/retrieve",
            json=body,
            timeout=self.request_timeout,
        )
        if response.status_code != 200:
            raise RuntimeError("Retrieval failed: %s" % response.text[:300])
        return response.json()

    def _retrieve_precedents_safe(
        self,
        tenant_id: str,
        query: str,
        top_k: int,
        version_mode: str,
        contract_type: Optional[str],
    ) -> Dict[str, Any]:
        try:
            return self._retrieve_precedents(
                tenant_id=tenant_id,
                query=query,
                top_k=top_k,
                version_mode=version_mode,
                contract_type=contract_type,
            )
        except Exception as exc:
            logger.warning("Precedent retrieval failed for query '%s': %s", query[:120], exc)
            return {"total": 0, "results": [], "error": str(exc)}

    def _extract_negotiation_intelligence(
        self,
        redline: RedlineContext,
        section_analysis: Dict[str, Any],
        retrieval_results: List[Dict[str, Any]],
    ) -> Dict[str, Any]:
        citations = []
        for item in retrieval_results[:6]:
            citations.append(
                {
                    "file": item.get("filename", ""),
                    "contract_type": item.get("contract_type", ""),
                    "score": item.get("score", 0),
                    "text": (item.get("text") or "")[:320],
                }
            )

        prompt = (
            "Extract negotiation intelligence and return JSON with keys: "
            "objection_pattern, our_response, final_outcome, success_factors, key_learnings.\n\n"
            "REDLINE: %s\nSECTION ANALYSIS: %s\nPRECEDENT CITATIONS: %s"
        ) % (
            redline.redline_text[:900],
            json_safe(section_analysis),
            json_safe(citations),
        )

        response = self.llm.query(
            prompt=prompt,
            system_message="You are a legal negotiation analyst. Return JSON only.",
            temperature=0.1,
            max_tokens=700,
        )
        if not response.success:
            return {
                "objection_pattern": "unknown",
                "our_response": "unknown",
                "final_outcome": "unknown",
                "success_factors": [],
                "key_learnings": "insufficient evidence",
            }

        try:
            return self.llm.parse_json(response.content)
        except Exception:
            return {
                "objection_pattern": "unknown",
                "our_response": "unknown",
                "final_outcome": "unknown",
                "success_factors": [],
                "key_learnings": "invalid llm response",
            }

    def _pick_best_query(self, queries: Dict[str, List[str]], fallback: str) -> str:
        for key in ["primary_queries", "secondary_queries", "comment_searches", "section_indicators"]:
            values = queries.get(key, [])
            if values:
                for candidate in values:
                    if (candidate or "").strip():
                        return candidate.strip()
        return fallback or "contract negotiation"

    def _build_surrounding_context(self, full_text: str, needle: str, window: int = 450) -> str:
        if not needle:
            return full_text[:window * 2]
        idx = full_text.find(needle)
        if idx < 0:
            return full_text[:window * 2]
        start = max(0, idx - window)
        end = min(len(full_text), idx + len(needle) + window)
        text = full_text[start:end]
        if start > 0:
            text = "..." + text
        if end < len(full_text):
            text = text + "..."
        return text

    def _build_baseline_query(self, full_text: str, comments: List[str]) -> str:
        chunks = []
        header = full_text[:900].strip()
        if header:
            chunks.append(header)
        if comments:
            chunks.append(" ".join(comments[:5])[:500])
        baseline = " ".join(chunks).strip()
        return baseline if baseline else "contract clause precedent analysis"

    def _md5(self, path: str) -> str:
        digest = hashlib.md5()
        with open(path, "rb") as f:
            for block in iter(lambda: f.read(8192), b""):
                digest.update(block)
        return digest.hexdigest()


def json_safe(value: Any) -> str:
    try:
        import json

        return json.dumps(value, ensure_ascii=False)
    except Exception:
        return str(value)
