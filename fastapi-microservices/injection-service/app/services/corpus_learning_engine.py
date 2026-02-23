import hashlib
import json
import logging
import os
import re
from datetime import datetime
from typing import Any, Callable, Dict, List, Optional, Tuple
from xml.etree import ElementTree as ET

from docx import Document

from .lm_studio_client import LMStudioClient
from .vector_store import PrecedentChunk, VectorStore

logger = logging.getLogger(__name__)

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


class CorpusLearningEngine:
    def __init__(
        self,
        progress_callback: Optional[Callable[[str], None]] = None,
        lm_studio_url: str = "http://localhost:1234",
        model: str = "mistral",
    ):
        self.progress_callback = progress_callback or (lambda msg: None)
        self.llm_client = LMStudioClient(base_url=lm_studio_url, model=model)

    def _report(self, message: str):
        logger.info(message)
        self.progress_callback(message)

    def learn_from_corpus(
        self,
        tenant_id: str,
        corpus_path: str,
        vector_store_path: str,
        upload_id: str,
        corpus_version: str,
    ) -> Dict[str, Any]:
        started_at = datetime.utcnow()
        self._report("Starting corpus learning for tenant '%s'" % tenant_id)

        vector_store = VectorStore(store_path=vector_store_path)
        vector_store.load_or_create()

        hash_file = os.path.join(vector_store_path, "processed_hashes_%s.json" % tenant_id)
        processed_hashes = self._read_json(hash_file, default={})
        version_file = os.path.join(vector_store_path, "document_versions_%s.json" % tenant_id)
        document_versions = self._read_json(version_file, default={})

        customer_dirs = self._discover_customer_dirs(corpus_path)
        stats = {
            "tenant_id": tenant_id,
            "upload_id": upload_id,
            "corpus_version": corpus_version,
            "total_customers_found": len(customer_dirs),
            "customers_analyzed": 0,
            "customers_skipped": 0,
            "total_documents": 0,
            "documents_skipped_unchanged": 0,
            "documents_failed": 0,
            "documents_versioned": 0,
            "chunks_deactivated": 0,
            "total_chunks_added": 0,
            "chunks_by_type": {"section": 0, "change": 0, "comment": 0},
            "agreement_type_counts": {},
            "clauses_learned": {},
            "errors": [],
            "started_at": started_at.isoformat(),
            "completed_at": None,
            "learning_time_seconds": 0,
            "index_stats": {},
        }

        for idx, customer_path in enumerate(customer_dirs):
            customer_name = os.path.basename(customer_path)
            self._report("[%s/%s] Processing customer: %s" % (idx + 1, len(customer_dirs), customer_name))

            try:
                customer_result = self._process_customer(
                    tenant_id=tenant_id,
                    customer_path=customer_path,
                    customer_name=customer_name,
                    processed_hashes=processed_hashes,
                    document_versions=document_versions,
                    vector_store=vector_store,
                    upload_id=upload_id,
                    corpus_version=corpus_version,
                )

                if customer_result["documents_total"] == 0:
                    stats["customers_skipped"] += 1
                else:
                    stats["customers_analyzed"] += 1

                stats["total_documents"] += customer_result["documents_processed"]
                stats["documents_skipped_unchanged"] += customer_result["documents_skipped_unchanged"]
                stats["documents_failed"] += customer_result["documents_failed"]
                stats["documents_versioned"] += customer_result["documents_versioned"]
                stats["chunks_deactivated"] += customer_result["chunks_deactivated"]
                stats["total_chunks_added"] += customer_result["chunks_added"]

                for key, value in customer_result["chunks_by_type"].items():
                    stats["chunks_by_type"][key] = stats["chunks_by_type"].get(key, 0) + value

                for key, value in customer_result["agreement_type_counts"].items():
                    stats["agreement_type_counts"][key] = stats["agreement_type_counts"].get(key, 0) + value

                for key, value in customer_result["clauses_learned"].items():
                    stats["clauses_learned"][key] = stats["clauses_learned"].get(key, 0) + value

            except Exception as exc:
                err = "%s: %s" % (customer_name, exc)
                stats["errors"].append(err)
                logger.exception("Customer processing failed: %s", customer_name)

        vector_store.save()
        self._write_json(hash_file, processed_hashes)
        self._write_json(version_file, document_versions)

        completed_at = datetime.utcnow()
        stats["completed_at"] = completed_at.isoformat()
        stats["learning_time_seconds"] = (completed_at - started_at).total_seconds()
        stats["index_stats"] = vector_store.get_stats()

        summary_path = os.path.join(vector_store_path, "precedent_index_%s.json" % tenant_id)
        self._write_json(summary_path, stats)

        self._report(
            "Completed learning. customers=%s documents=%s chunks=%s"
            % (stats["customers_analyzed"], stats["total_documents"], stats["total_chunks_added"])
        )
        return stats

    def _discover_customer_dirs(self, corpus_path: str) -> List[str]:
        if not os.path.exists(corpus_path):
            raise FileNotFoundError("Corpus path not found: %s" % corpus_path)

        dirs = []
        for item in sorted(os.listdir(corpus_path)):
            full = os.path.join(corpus_path, item)
            if os.path.isdir(full) and not item.startswith((".", "_")):
                dirs.append(full)
        return dirs

    def _process_customer(
        self,
        tenant_id: str,
        customer_path: str,
        customer_name: str,
        processed_hashes: Dict[str, Any],
        document_versions: Dict[str, Any],
        vector_store: VectorStore,
        upload_id: str,
        corpus_version: str,
    ) -> Dict[str, Any]:
        result = {
            "documents_total": 0,
            "documents_processed": 0,
            "documents_skipped_unchanged": 0,
            "documents_failed": 0,
            "documents_versioned": 0,
            "chunks_deactivated": 0,
            "chunks_added": 0,
            "chunks_by_type": {"section": 0, "change": 0, "comment": 0},
            "agreement_type_counts": {},
            "clauses_learned": {},
        }

        docx_files = self._list_docx_recursive(customer_path)
        result["documents_total"] = len(docx_files)
        if not docx_files:
            return result

        for seq_order, rel_path in enumerate(docx_files, start=1):
            file_path = os.path.join(customer_path, rel_path)
            hash_key = "%s/%s/%s" % (tenant_id, customer_name, rel_path)
            document_id = hash_key
            file_hash = self._md5(file_path)

            old_hash_entry = processed_hashes.get(hash_key)
            previous_hash = old_hash_entry
            if isinstance(old_hash_entry, dict):
                previous_hash = old_hash_entry.get("file_hash")

            if previous_hash == file_hash:
                result["documents_skipped_unchanged"] += 1
                continue

            try:
                extraction = self._extract_document(file_path)
                contract_type = self.llm_client.classify_contract_type(
                    extraction["content"], rel_path
                )

                previous_document_version = int(document_versions.get(document_id, 0))
                next_document_version = previous_document_version + 1
                if previous_document_version > 0:
                    result["documents_versioned"] += 1
                    deactivated = vector_store.deactivate_document(
                        tenant_id=tenant_id,
                        document_id=document_id,
                    )
                    result["chunks_deactivated"] += deactivated

                chunks = self._create_chunks(
                    tenant_id=tenant_id,
                    upload_id=upload_id,
                    corpus_version=corpus_version,
                    document_id=document_id,
                    document_version=next_document_version,
                    customer=customer_name,
                    filename=rel_path,
                    source_path=file_path,
                    contract_type=contract_type,
                    seq_order=seq_order,
                    extraction=extraction,
                    doc_hash=file_hash,
                )

                added = vector_store.add_chunks(chunks)
                result["chunks_added"] += added
                result["documents_processed"] += 1
                document_versions[document_id] = next_document_version
                processed_hashes[hash_key] = {
                    "file_hash": file_hash,
                    "document_version": next_document_version,
                    "last_upload_id": upload_id,
                    "last_updated_at": datetime.utcnow().isoformat(),
                }

                result["agreement_type_counts"][contract_type] = (
                    result["agreement_type_counts"].get(contract_type, 0) + 1
                )

                for chunk in chunks:
                    chunk_type = chunk.chunk_type
                    result["chunks_by_type"][chunk_type] = result["chunks_by_type"].get(chunk_type, 0) + 1
                    if chunk.clause_label:
                        result["clauses_learned"][chunk.clause_label] = (
                            result["clauses_learned"].get(chunk.clause_label, 0) + 1
                        )

            except Exception:
                result["documents_failed"] += 1
                logger.exception("Document failed: %s", file_path)

        return result

    def _list_docx_recursive(self, customer_path: str) -> List[str]:
        files = []  # type: List[str]
        for root, _, filenames in os.walk(customer_path):
            for name in filenames:
                if name.startswith("~"):
                    continue
                if not name.lower().endswith(".docx"):
                    continue
                abs_path = os.path.join(root, name)
                rel_path = os.path.relpath(abs_path, customer_path)
                files.append(rel_path.replace("\\", "/"))
        files.sort()
        return files

    def _extract_document(self, file_path: str) -> Dict[str, Any]:
        doc = Document(file_path)
        text_blocks = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        content = "\n".join(text_blocks)

        changes, confidence = self._extract_tracked_changes(doc)
        comments = self._extract_comments(doc)

        return {
            "content": content,
            "tracked_changes": changes,
            "comments": comments,
            "extraction_confidence": confidence,
        }

    def _extract_tracked_changes(self, doc: Document) -> Tuple[List[Dict[str, str]], float]:
        changes = []  # type: List[Dict[str, str]]
        confidence = 0.3

        try:
            xml = doc.element.xml
            root = ET.fromstring(xml)

            for el in root.iter():
                tag = el.tag
                if tag.endswith("}ins"):
                    text = self._collect_text(el)
                    if text:
                        changes.append({"type": "ADDITION", "text": text})
                elif tag.endswith("}del"):
                    text = self._collect_text(el)
                    if text:
                        changes.append({"type": "DELETION", "text": text})

            if changes:
                confidence = 0.9
        except Exception:
            logger.exception("XML tracked-change parsing failed; using run-style fallback")

        if not changes:
            # Weak fallback when true tracked-change data is unavailable.
            for para in doc.paragraphs:
                for run in para.runs:
                    txt = (run.text or "").strip()
                    if not txt:
                        continue
                    if run.font.strike:
                        changes.append({"type": "STRIKEOUT", "text": txt})
                    elif run.underline:
                        changes.append({"type": "UNDERLINE", "text": txt})
            if changes:
                confidence = 0.5

        return changes, confidence

    def _extract_comments(self, doc: Document) -> List[str]:
        comments = []
        markers = ["comment:", "note:", "rationale:", "requested:", "proposed:", "discussed"]

        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if not text:
                continue
            lowered = text.lower()
            if any(marker in lowered for marker in markers):
                comments.append(text)

        return comments

    def _create_chunks(
        self,
        tenant_id: str,
        upload_id: str,
        corpus_version: str,
        document_id: str,
        document_version: int,
        customer: str,
        filename: str,
        source_path: str,
        contract_type: str,
        seq_order: int,
        extraction: Dict[str, Any],
        doc_hash: str,
    ) -> List[PrecedentChunk]:
        chunks = []  # type: List[PrecedentChunk]

        sections = self._split_into_clause_sections(extraction["content"])
        for clause_label, section_text in sections:
            for text in self._chunk_text(section_text, chunk_size=700, overlap=120):
                chunk_hash = self._hash_chunk(doc_hash, text, "section", clause_label)
                chunks.append(
                    PrecedentChunk(
                        text=text,
                        tenant_id=tenant_id,
                        upload_id=upload_id,
                        corpus_version=corpus_version,
                        document_id=document_id,
                        document_version=document_version,
                        customer=customer,
                        filename=filename,
                        source_path=source_path,
                        contract_type=contract_type,
                        chunk_type="section",
                        clause_label=clause_label,
                        sequence_order=seq_order,
                        extraction_confidence=extraction["extraction_confidence"],
                        doc_hash=doc_hash,
                        chunk_hash=chunk_hash,
                    )
                )

        for change in extraction["tracked_changes"]:
            text = (change.get("text") or "").strip()
            if len(text) < 8:
                continue
            chunk_hash = self._hash_chunk(doc_hash, text, "change", change.get("type", ""))
            chunks.append(
                PrecedentChunk(
                    text=text,
                    tenant_id=tenant_id,
                    upload_id=upload_id,
                    corpus_version=corpus_version,
                    document_id=document_id,
                    document_version=document_version,
                    customer=customer,
                    filename=filename,
                    source_path=source_path,
                    contract_type=contract_type,
                    chunk_type="change",
                    clause_label="redline",
                    change_type=change.get("type", ""),
                    sequence_order=seq_order,
                    extraction_confidence=extraction["extraction_confidence"],
                    doc_hash=doc_hash,
                    chunk_hash=chunk_hash,
                )
            )

        for comment in extraction["comments"]:
            chunk_hash = self._hash_chunk(doc_hash, comment, "comment", "comment")
            chunks.append(
                PrecedentChunk(
                    text=comment,
                    tenant_id=tenant_id,
                    upload_id=upload_id,
                    corpus_version=corpus_version,
                    document_id=document_id,
                    document_version=document_version,
                    customer=customer,
                    filename=filename,
                    source_path=source_path,
                    contract_type=contract_type,
                    chunk_type="comment",
                    clause_label="comment",
                    sequence_order=seq_order,
                    extraction_confidence=extraction["extraction_confidence"],
                    doc_hash=doc_hash,
                    chunk_hash=chunk_hash,
                )
            )

        return chunks

    def _split_into_clause_sections(self, content: str) -> List[Tuple[str, str]]:
        lines = [line.strip() for line in content.split("\n") if line.strip()]
        if not lines:
            return []

        heading_re = re.compile(r"^(\d+(\.\d+)*)\s+([A-Z][A-Za-z\s\-/]{2,80})$")
        sections = []  # type: List[Tuple[str, str]]

        current_heading = "general"
        buffer = []  # type: List[str]

        for line in lines:
            match = heading_re.match(line)
            if match and len(buffer) >= 1:
                sections.append((current_heading, "\n".join(buffer)))
                current_heading = match.group(3).strip().lower()
                buffer = [line]
            else:
                buffer.append(line)

        if buffer:
            sections.append((current_heading, "\n".join(buffer)))

        return sections

    def _chunk_text(self, text: str, chunk_size: int, overlap: int) -> List[str]:
        text = text.strip()
        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0
        while start < len(text):
            end = min(start + chunk_size, len(text))
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            if end == len(text):
                break
            start = max(0, end - overlap)
        return chunks

    def _collect_text(self, xml_node: ET.Element) -> str:
        tokens = []
        for t in xml_node.iter():
            if t.tag.endswith("}t") or t.tag.endswith("}delText"):
                val = (t.text or "").strip()
                if val:
                    tokens.append(val)
        return " ".join(tokens).strip()

    def _md5(self, file_path: str) -> str:
        digest = hashlib.md5()
        with open(file_path, "rb") as f:
            for block in iter(lambda: f.read(8192), b""):
                digest.update(block)
        return digest.hexdigest()

    def _hash_chunk(self, doc_hash: str, text: str, chunk_type: str, label: str) -> str:
        base = "%s|%s|%s|%s" % (doc_hash, chunk_type, label, text)
        return hashlib.sha256(base.encode("utf-8")).hexdigest()

    def _read_json(self, path: str, default: Any) -> Any:
        if not os.path.exists(path):
            return default
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)

    def _write_json(self, path: str, payload: Any):
        os.makedirs(os.path.dirname(path), exist_ok=True)
        tmp_path = path + ".tmp"
        with open(tmp_path, "w", encoding="utf-8") as f:
            json.dump(payload, f, indent=2, ensure_ascii=False)
        os.replace(tmp_path, path)
